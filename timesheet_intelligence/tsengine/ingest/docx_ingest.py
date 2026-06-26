"""DOCX extractor.

Reads paragraphs and tables via python-docx. Crucially, many "Word" timesheets
are just a screenshot pasted into a document (no text, no tables) -- so when the
body is empty we extract the embedded media images and route them through the
image/OCR/vision path instead.
"""
from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Optional

from ..schema import (ExtractionQuality, FileKind, RawExtraction, RawImage,
                      RawTable, SourceRef)
from ..settings import Settings, get_settings
from .excel import _name_hint
from .ocr import layout_ocr, tesseract_available


def extract(path: str | Path, settings: Optional[Settings] = None) -> RawExtraction:
    s = settings or get_settings()
    p = Path(path)
    rel = p.name
    raw = RawExtraction(file=rel, kind=FileKind.DOCX, quality=ExtractionQuality.NATIVE)
    raw.meta["name_hint"] = _name_hint(p.stem)

    text_lines: list[str] = []
    try:
        import docx

        d = docx.Document(str(p))
        for para in d.paragraphs:
            if para.text.strip():
                text_lines.append(para.text)
        for ti, t in enumerate(d.tables):
            rows = [[c.text.strip() for c in r.cells] for r in t.rows]
            if rows:
                headers = rows[0]
                raw.tables.append(RawTable(
                    headers=headers, rows=rows[1:], title=f"table {ti + 1}",
                    source=SourceRef(file=rel, extractor="docx"),
                ))
                for r in rows:
                    text_lines.append(" | ".join(r))
    except Exception as exc:
        raw.notes.append(f"python-docx failed: {exc}")

    raw.text = "\n".join(text_lines).strip()
    raw.sources.append(SourceRef(file=rel, extractor="docx"))

    # Pull embedded media ALWAYS (not just when the body is empty): many Word
    # timesheets paste the hours as images while still carrying some text (week
    # labels, a header), which previously hid the images from the vision path.
    media = _extract_media(p, s)
    if media:
        ocr_parts: list[str] = []
        have_ocr = tesseract_available(s) and s.use_local_ocr
        for idx, img_path in enumerate(media, start=1):
            from PIL import Image

            try:
                with Image.open(img_path) as im:
                    w, h = im.size
            except Exception:
                w = h = None
            raw.images.append(RawImage(
                path=img_path, width=w, height=h,
                source=SourceRef(file=rel, region=f"embedded image {idx}",
                                 extractor="docx_media"),
            ))
            if have_ocr:
                txt, conf = layout_ocr(img_path, s)
                if txt.strip():
                    ocr_parts.append(f"----- embedded image {idx} (OCR) -----\n{txt}")
                    raw.meta["ocr_confidence"] = conf
        if ocr_parts:
            raw.text = (raw.text + "\n" + "\n".join(ocr_parts)).strip()
        # No real table parsed -> the hours live in the images: mark for vision.
        if not raw.tables:
            raw.quality = ExtractionQuality.OCR if ocr_parts else ExtractionQuality.VISION
            raw.notes.append("DOCX hours are in embedded images; OCR/vision used")
    elif not raw.text and not raw.tables:
        raw.quality = ExtractionQuality.EMPTY
        raw.notes.append("empty docx with no embedded media")

    return raw


def _extract_media(p: Path, s: Settings) -> list[str]:
    out_dir = s.output_path / "_evidence" / p.stem
    out_dir.mkdir(parents=True, exist_ok=True)
    paths: list[str] = []
    MAX_ENTRY = 50 * 1024 * 1024      # 50 MB per media file
    MAX_FILES = 30                    # cap number of embedded images
    try:
        with zipfile.ZipFile(p) as z:
            for info in z.infolist():
                name = info.filename
                low = name.lower()
                if not (low.startswith("word/media/") and low.split(".")[-1] in (
                        "png", "jpg", "jpeg", "tif", "tiff", "bmp", "gif", "webp")):
                    continue
                if info.file_size > MAX_ENTRY or len(paths) >= MAX_FILES:
                    continue          # skip oversized / excess media (zip-bomb guard)
                target = out_dir / Path(name).name
                with z.open(info) as src, open(target, "wb") as dst:
                    # bounded streaming copy
                    remaining = MAX_ENTRY
                    while remaining > 0:
                        chunk = src.read(1 << 20)
                        if not chunk:
                            break
                        dst.write(chunk)
                        remaining -= len(chunk)
                paths.append(str(target))
    except Exception:
        pass
    return paths
